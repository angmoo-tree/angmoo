from __future__ import annotations

import asyncio
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from io import BytesIO
import hashlib
import math
import multiprocessing
import os
from pathlib import PurePosixPath
import re
import stat
import sys
import time
from typing import Any, Callable
from uuid import uuid4
from zipfile import BadZipFile, ZipFile
from zoneinfo import ZoneInfo

from pypdf import PdfReader
from sqlalchemy import func, select, update
from sqlalchemy.orm import Session

from app import models, schemas
from app.cruds import agents as agent_crud
from app.cruds import community as community_crud
from app.credentials import (
    CredentialPurpose,
    CredentialResolutionError,
    CredentialResolver,
)
from app.providers.contracts import EmbeddingRequest
from app.providers.registry import get_embedding_adapter
from app.services.direct_llm import (
    DirectLlmCallContext,
    RunLlmTracker,
    wait_for_provider_rate_limit,
)
from app.services import community as community_service
from app.services import lore_parser_quota
from app.core.context_text import neutralize_context_text


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_LORE_FILE_BYTES = 10 * 1024 * 1024
LORE_UPLOAD_READ_CHUNK_BYTES = 64 * 1024
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MAX_DOCX_ENTRIES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_XML_BYTES = 20 * 1024 * 1024
MAX_DOCX_ENTRY_BYTES = 10 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_PDF_PAGES = 200
MAX_PDF_OBJECTS = 20_000
MAX_PARSED_TEXT_CHARS = 100_000
# Windows uses the spawn multiprocessing start method and must import the
# application module graph in the isolated child.  Real-time antivirus scanning
# can make that startup alone exceed the POSIX parser budget even for a tiny,
# valid document.  Keep the parser isolated and bounded while allowing the
# measured Windows startup overhead.
LORE_PARSER_TIMEOUT_SECONDS = 15.0 if os.name == "nt" else 8.0
LORE_PARSER_MEMORY_BYTES = 512 * 1024 * 1024
LORE_PARSER_CPU_SECONDS = 6
MAX_LORE_SOURCES_PER_CHARACTER = 1
MAX_LORE_TEXT_CHARS_PER_CHARACTER = 50_000
MAX_LORE_CHUNKS_PER_CHARACTER = 100
TARGET_CHUNK_MIN_CHARS = 500
TARGET_CHUNK_TARGET_CHARS = 750
TARGET_CHUNK_MAX_CHARS = 1000
EMBEDDING_MODEL = "gemini-embedding-2"
EMBEDDING_DIMENSION = 768
RETRIEVAL_CANDIDATE_LIMIT = 30
RETRIEVAL_FINAL_LIMIT = 5
RECENT_LORE_USAGE_LIMIT = 12
RECENT_LORE_STRONG_PENALTY_WINDOW = timedelta(days=2)
RECENT_LORE_SOFT_PENALTY_WINDOW = timedelta(days=7)
APP_TIMEZONE = ZoneInfo("Asia/Seoul")


class CharacterLoreError(Exception):
    pass


class CharacterLoreNotFoundError(CharacterLoreError):
    pass


class CharacterLoreValidationError(CharacterLoreError):
    pass


class CharacterLoreFileTooLargeError(CharacterLoreValidationError):
    pass


class CharacterLoreEmbeddingError(CharacterLoreError):
    pass


class CharacterLoreParserBusyError(CharacterLoreError):
    retry_after_seconds = lore_parser_quota.RETRY_AFTER_SECONDS


@dataclass(frozen=True)
class LoreChunkDraft:
    section_hint: str | None
    text: str
    content_hash: str


@dataclass(frozen=True)
class RetrievedLoreChunk:
    id: str
    source_id: str
    source_filename: str
    section_hint: str | None
    text: str
    distance: float


@dataclass(frozen=True)
class LoreRetrievalResult:
    mode: str
    chunks: tuple[RetrievedLoreChunk, ...] = ()
    error_message: str | None = None

    @property
    def chunk_ids(self) -> list[str]:
        return [chunk.id for chunk in self.chunks]


@dataclass(frozen=True)
class _GoogleEmbeddingCredential:
    api_key: str
    credential_id: str
    key_fingerprint: str | None
    provider: str
    model: str = EMBEDDING_MODEL


async def read_lore_upload_bytes(upload_file: Any) -> bytes:
    content = bytearray()
    while True:
        remaining = MAX_LORE_FILE_BYTES + 1 - len(content)
        if remaining <= 0:
            raise CharacterLoreFileTooLargeError(
                "설정집 파일은 10 MiB 이하만 업로드할 수 있습니다."
            )
        chunk = await upload_file.read(min(LORE_UPLOAD_READ_CHUNK_BYTES, remaining))
        if not chunk:
            return bytes(content)
        content.extend(chunk)
        if len(content) > MAX_LORE_FILE_BYTES:
            raise CharacterLoreFileTooLargeError(
                "설정집 파일은 10 MiB 이하만 업로드할 수 있습니다."
            )


def list_lore_sources(
    db: Session, user: models.User, character_id: str
) -> list[schemas.CharacterLoreSourceRead]:
    _get_owned_character(db, user, character_id)
    sources = list(
        db.scalars(
            select(models.CharacterLoreSource)
            .where(models.CharacterLoreSource.character_id == character_id)
            .order_by(
                models.CharacterLoreSource.created_at.desc(),
                models.CharacterLoreSource.id.desc(),
            )
        )
    )
    return [schemas.CharacterLoreSourceRead.model_validate(source) for source in sources]


def lore_status(
    db: Session, user: models.User, character_id: str
) -> schemas.CharacterLoreStatusRead:
    _get_owned_character(db, user, character_id)
    return _status_read(db, character_id)


def upload_lore_source(
    db: Session,
    user: models.User,
    character_id: str,
    *,
    filename: str,
    content_type: str | None,
    file_bytes: bytes,
    replace_existing: bool = False,
) -> schemas.CharacterLoreSourceRead:
    extension = _validated_extension(filename)
    if len(file_bytes) > MAX_LORE_FILE_BYTES:
        raise CharacterLoreFileTooLargeError(
            "설정집 파일은 10 MiB 이하만 업로드할 수 있습니다."
        )
    validate_lore_upload_contract(
        extension=extension,
        content_type=content_type,
        file_bytes=file_bytes,
    )
    character = _get_owned_character(db, user, character_id)

    if extension in {".pdf", ".docx"}:
        try:
            with lore_parser_quota.parser_lease(db, user_id=user.id):
                raw_text = _extract_text(extension, file_bytes)
        except (
            lore_parser_quota.LoreParserCapacityError,
            lore_parser_quota.LoreParserLeaseUnavailableError,
        ) as exc:
            raise CharacterLoreParserBusyError(
                "Document parser capacity is temporarily unavailable"
            ) from exc
    else:
        raw_text = _extract_text(extension, file_bytes)
    raw_hash = _sha256_text(raw_text)
    existing = db.scalar(
        select(models.CharacterLoreSource)
        .where(
            models.CharacterLoreSource.character_id == character.id,
            models.CharacterLoreSource.raw_text_hash == raw_hash,
        )
        .limit(1)
    )
    if existing is not None:
        return schemas.CharacterLoreSourceRead.model_validate(existing)
    existing_sources = list(
        db.scalars(
            select(models.CharacterLoreSource).where(
                models.CharacterLoreSource.character_id == character.id
            )
        )
    )
    if existing_sources and not replace_existing:
        raise CharacterLoreValidationError(
            "앵무 1마리당 설정집 파일은 1개만 업로드할 수 있습니다. 기존 설정집을 삭제하거나 교체해 주세요."
        )

    drafts = chunk_lore_text(raw_text)
    replacing_text_chars = (
        sum(source.extracted_char_count for source in existing_sources)
        if replace_existing
        else 0
    )
    replacing_chunk_count = (
        sum(source.chunk_count for source in existing_sources) if replace_existing else 0
    )
    _ensure_character_limits(
        db,
        character_id=character.id,
        new_text_chars=len(raw_text),
        new_chunk_count=len(drafts),
        replacing_text_chars=replacing_text_chars,
        replacing_chunk_count=replacing_chunk_count,
    )
    for existing_source in existing_sources:
        db.delete(existing_source)
    source = models.CharacterLoreSource(
        id=f"lore-src-{uuid4().hex[:12]}",
        owner_id=user.id,
        character_id=character.id,
        filename=_safe_filename(filename),
        extension=extension.lstrip("."),
        content_type=(content_type or "")[:120] or None,
        file_size_bytes=len(file_bytes),
        raw_text=raw_text,
        raw_text_hash=raw_hash,
        extracted_char_count=len(raw_text),
        chunk_count=len(drafts),
        status="embedding_failed",
        error_message=None,
    )
    db.add(source)
    _store_chunks_with_embeddings(db, character=character, source=source, drafts=drafts)
    db.commit()
    db.refresh(source)
    return schemas.CharacterLoreSourceRead.model_validate(source)


def delete_lore_source(
    db: Session, user: models.User, character_id: str, source_id: str
) -> None:
    _get_owned_character(db, user, character_id)
    source = _get_owned_source(db, character_id=character_id, source_id=source_id)
    db.delete(source)
    db.commit()


def rebuild_lore_source(
    db: Session, user: models.User, character_id: str, source_id: str
) -> schemas.CharacterLoreSourceRead:
    character = _get_owned_character(db, user, character_id)
    source = _get_owned_source(db, character_id=character.id, source_id=source_id)
    drafts = chunk_lore_text(source.raw_text)
    existing_text_chars = _total_text_chars(db, character.id) - source.extracted_char_count
    existing_chunk_count = _chunk_count(db, character.id) - source.chunk_count
    if existing_text_chars + len(source.raw_text) > MAX_LORE_TEXT_CHARS_PER_CHARACTER:
        raise CharacterLoreValidationError("설정집 원문은 앵무당 최대 50,000자까지 가능합니다.")
    if existing_chunk_count + len(drafts) > MAX_LORE_CHUNKS_PER_CHARACTER:
        raise CharacterLoreValidationError("설정집 chunk는 앵무당 최대 100개까지 가능합니다.")
    reusable = {
        chunk.content_hash: chunk.embedding
        for chunk in source.chunks
        if chunk.status == "ready"
        and chunk.embedding is not None
        and chunk.embedding_model == EMBEDDING_MODEL
        and chunk.embedding_dimension == EMBEDDING_DIMENSION
    }
    for chunk in list(source.chunks):
        db.delete(chunk)
    db.flush()
    source.chunk_count = len(drafts)
    source.status = "embedding_failed"
    source.error_message = None
    _store_chunks_with_embeddings(
        db, character=character, source=source, drafts=drafts, reusable_embeddings=reusable
    )
    db.commit()
    db.refresh(source)
    return schemas.CharacterLoreSourceRead.model_validate(source)


def chunk_lore_text(text: str) -> list[LoreChunkDraft]:
    normalized = _normalize_text(text)
    if not normalized:
        raise CharacterLoreValidationError("설정집에서 텍스트를 추출하지 못했습니다.")
    units = _split_lore_units(normalized)
    chunks: list[tuple[str | None, str]] = []
    current_section: str | None = None
    current_parts: list[str] = []
    current_length = 0

    def flush_current() -> None:
        nonlocal current_parts, current_length, current_section
        if not current_parts:
            return
        chunks.append((current_section, "\n".join(current_parts).strip()))
        current_parts = []
        current_length = 0

    for section_hint, unit in units:
        for piece in _split_long_unit(unit):
            piece_length = len(piece)
            if current_parts:
                next_length = current_length + piece_length + 1
                if section_hint != current_section or next_length > TARGET_CHUNK_MAX_CHARS:
                    flush_current()
            if not current_parts:
                current_section = section_hint
            current_parts.append(piece)
            current_length += piece_length + 1
            if current_length >= TARGET_CHUNK_TARGET_CHARS:
                flush_current()
    flush_current()

    drafts = [
        LoreChunkDraft(
            section_hint=section,
            text=chunk_text,
            content_hash=_sha256_text(chunk_text),
        )
        for section, chunk_text in chunks
        if chunk_text
    ]
    if not drafts:
        raise CharacterLoreValidationError("설정집에서 검색 가능한 chunk를 만들지 못했습니다.")
    return drafts


def retrieve_lore_for_self_update(
    db: Session,
    *,
    character: models.Character,
    now: datetime | None = None,
) -> LoreRetrievalResult:
    query = build_lore_search_query(db, character=character, now=now)
    return retrieve_lore_for_query(db, character=character, query=query)


def has_ready_lore_chunks(db: Session, *, character_id: str) -> bool:
    return _ready_chunk_count(db, character_id) > 0


def retrieve_lore_for_query(
    db: Session,
    *,
    character: models.Character,
    query: str,
) -> LoreRetrievalResult:
    clean_query = query.strip()
    if not clean_query or _ready_chunk_count(db, character.id) <= 0:
        return LoreRetrievalResult(mode="fallback_no_lore")
    try:
        credential = _google_embedding_credential_for_character(db, character.id)
        query_embedding = _embed_text(
            credential.api_key, _query_embedding_input(clean_query)
        )
        rows = _retrieve_lore_rows(
            db, character_id=character.id, query_embedding=query_embedding
        )
    except Exception as exc:
        return LoreRetrievalResult(
            mode="fallback_embedding_failed",
            error_message=str(exc)[:500],
        )
    return _lore_result_from_rows(rows)


async def retrieve_lore_for_query_tracked(
    db: Session,
    *,
    character: models.Character,
    query: str,
    tracker: RunLlmTracker,
    agent_run_id: str,
    node: str = "CharacterLoreEmbedding",
    lane: str = "lore_query_embedding",
) -> LoreRetrievalResult:
    clean_query = query.strip()
    if not clean_query or _ready_chunk_count(db, character.id) <= 0:
        return LoreRetrievalResult(mode="fallback_no_lore")
    try:
        credential = _google_embedding_credential_for_character(db, character.id)
        context = DirectLlmCallContext(
            credential_id=credential.credential_id,
            key_fingerprint=credential.key_fingerprint,
            character_id=character.id,
            agent_run_id=agent_run_id,
            node=node,
            lane=lane,
            provider=credential.provider,
            model=credential.model,
        )
        query_embedding = await _embed_text_tracked(
            credential.api_key,
            _query_embedding_input(clean_query),
            context=context,
            tracker=tracker,
        )
        rows = _retrieve_lore_rows(
            db, character_id=character.id, query_embedding=query_embedding
        )
    except Exception as exc:
        return LoreRetrievalResult(
            mode="fallback_embedding_failed",
            error_message=str(exc)[:500],
        )
    return _lore_result_from_rows(rows)


def _retrieve_lore_rows(
    db: Session, *, character_id: str, query_embedding: list[float]
) -> list[tuple[models.CharacterLoreChunk, float]]:
    # A character is bounded to 100 lore chunks, so in-process cosine ranking
    # is deterministic and comfortably small.  This preserves retrieval
    # behavior without retaining a PostgreSQL/pgvector runtime dependency.
    chunks = list(
        db.scalars(
            select(models.CharacterLoreChunk)
            .join(models.CharacterLoreSource)
            .where(
                models.CharacterLoreChunk.character_id == character_id,
                models.CharacterLoreChunk.status == "ready",
                models.CharacterLoreChunk.embedding.is_not(None),
                models.CharacterLoreSource.status.in_(("ready", "partial")),
            )
            .order_by(models.CharacterLoreChunk.id)
        )
    )
    ranked = [
        (chunk, _cosine_distance(chunk.embedding or [], query_embedding))
        for chunk in chunks
    ]
    ranked.sort(key=lambda item: (item[1], item[0].id))
    return ranked[:RETRIEVAL_CANDIDATE_LIMIT]


def _cosine_distance(left: list[float], right: list[float]) -> float:
    if not left or len(left) != len(right):
        return 1.0
    dot = sum(a * b for a, b in zip(left, right, strict=True))
    left_norm = math.sqrt(sum(value * value for value in left))
    right_norm = math.sqrt(sum(value * value for value in right))
    if left_norm == 0.0 or right_norm == 0.0:
        return 1.0
    similarity = max(-1.0, min(1.0, dot / (left_norm * right_norm)))
    return 1.0 - similarity


def _lore_result_from_rows(
    rows: list[tuple[models.CharacterLoreChunk, float]],
) -> LoreRetrievalResult:
    selected = _rerank_lore_candidates(rows)
    if not selected:
        return LoreRetrievalResult(mode="fallback_no_lore")
    return LoreRetrievalResult(mode="pgvector", chunks=tuple(selected))


def build_lore_search_query(
    db: Session,
    *,
    character: models.Character,
    now: datetime | None = None,
) -> str:
    current_time = now or datetime.now(UTC)
    recent_topics = community_service.format_recent_own_root_topic_history_for_prompt(
        db, character_id=character.id
    )
    recent_lore = _format_recent_lore_usage(db, character_id=character.id)
    return "\n".join(
        [
            "이 캐릭터가 특정 커뮤니티 글에 답하는 것이 아니라 독립 root 글로 쓸 만한 내부 소재를 찾는다.",
            f"현재 KST 시간대 참고: {current_time.astimezone(APP_TIMEZONE).isoformat()}",
            f"캐릭터 이름: {character.name}",
            f"기본 페르소나: {character.persona_summary or '-'}",
            f"성격: {character.personality or '-'}",
            f"말투: {character.speech_style or '-'}",
            f"세계관/배경: {character.worldview or '-'}",
            f"관심 주제: {character.topic_preferences or '-'}",
            "최근 자기 root 글 topic 이력:",
            recent_topics,
            "최근 사용한 설정집 소재:",
            recent_lore,
            "찾을 자료: 자기 생각, 기억, 취향, 습관, 장소, 물건, 고민, 관찰, 세계관 안의 작은 소재.",
        ]
    )


def format_lore_prompt_context(
    result: LoreRetrievalResult,
    *,
    lore_query_mode: str | None = None,
    max_chunks: int | None = None,
    max_text_chars: int = 900,
) -> str:
    if not result.chunks:
        return ""
    chunks = result.chunks[:max_chunks] if max_chunks is not None else result.chunks
    lines = [
        "Character lore retrieval:",
        f"- retrieval_mode: {result.mode}",
        f"- lore_query_mode: {lore_query_mode}" if lore_query_mode else "",
        "- These chunks are private reference material. Do not copy sentences from them into title/body.",
        "- Use them only to choose a character-owned thought, memory, habit, taste, object, place, or worldview detail.",
        "- Do not write as if replying to a specific community post or author.",
        "- Do not expose lore_chunk_ids, retrieval_mode, lore_query_mode, or source filenames in visible title/body.",
        "selected_lore_chunks:",
    ]
    lines = [line for line in lines if line]
    for chunk in chunks:
        lines.extend(
            [
                f"- lore_chunk_id: {chunk.id}",
                f"  source: {neutralize_context_text(chunk.source_filename)}",
                f"  section_hint: {neutralize_context_text(chunk.section_hint or '-')}",
                "  text: " + neutralize_context_text(chunk.text)[:max_text_chars],
            ]
        )
    return "\n".join(lines)


def mark_lore_chunks_used(db: Session, *, chunk_ids: list[str]) -> None:
    if not chunk_ids:
        return
    db.execute(
        update(models.CharacterLoreChunk)
        .where(models.CharacterLoreChunk.id.in_(chunk_ids))
        .values(
            last_used_at=datetime.now(UTC),
            usage_count=models.CharacterLoreChunk.usage_count + 1,
        )
    )
    db.commit()


def _store_chunks_with_embeddings(
    db: Session,
    *,
    character: models.Character,
    source: models.CharacterLoreSource,
    drafts: list[LoreChunkDraft],
    reusable_embeddings: dict[str, list[float]] | None = None,
) -> None:
    reusable = dict(reusable_embeddings or {})
    reusable.update(_existing_embeddings_by_hash(db, character.id))
    api_key = ""
    embedding_error: str | None = None
    try:
        api_key = _google_api_key_for_character(db, character.id)
    except CharacterLoreEmbeddingError as exc:
        embedding_error = str(exc)

    ready_count = 0
    for index, draft in enumerate(drafts):
        embedding = reusable.get(draft.content_hash)
        chunk_error = embedding_error
        if embedding is None and api_key:
            try:
                embedding = _embed_text(api_key, _chunk_embedding_input(draft))
                chunk_error = None
            except Exception as exc:
                chunk_error = str(exc)[:500]
        status = "ready" if embedding is not None else "embedding_failed"
        if status == "ready":
            ready_count += 1
        db.add(
            models.CharacterLoreChunk(
                id=f"lore-chunk-{uuid4().hex[:12]}",
                source_id=source.id,
                owner_id=character.owner_id,
                character_id=character.id,
                chunk_index=index,
                section_hint=draft.section_hint,
                text=draft.text,
                content_hash=draft.content_hash,
                embedding=embedding,
                embedding_model=EMBEDDING_MODEL if embedding is not None else None,
                embedding_dimension=EMBEDDING_DIMENSION if embedding is not None else None,
                status=status,
                error_message=chunk_error,
                usage_count=0,
            )
        )
    if ready_count == len(drafts):
        source.status = "ready"
        source.error_message = None
    elif ready_count > 0:
        source.status = "partial"
        source.error_message = "일부 chunk embedding에 실패했습니다."
    else:
        source.status = "embedding_failed"
        source.error_message = embedding_error or "설정집 embedding에 실패했습니다."


def _get_owned_character(
    db: Session, user: models.User, character_id: str
) -> models.Character:
    character = community_crud.get_character(db, character_id)
    if character is None or character.deleted_at is not None or character.owner_id != user.id:
        raise CharacterLoreNotFoundError(character_id)
    return character


def _get_owned_source(
    db: Session, *, character_id: str, source_id: str
) -> models.CharacterLoreSource:
    source = db.get(models.CharacterLoreSource, source_id)
    if source is None or source.character_id != character_id:
        raise CharacterLoreNotFoundError(source_id)
    return source


def _validated_extension(filename: str) -> str:
    lower = filename.strip().lower()
    if lower.endswith(".doc") and not lower.endswith(".docx"):
        raise CharacterLoreValidationError("구형 .doc 파일은 v1에서 지원하지 않습니다.")
    extension = "." + lower.rsplit(".", 1)[-1] if "." in lower else ""
    if extension not in SUPPORTED_EXTENSIONS:
        raise CharacterLoreValidationError("PDF, DOCX, TXT, MD 파일만 업로드할 수 있습니다.")
    return extension


def _safe_filename(filename: str) -> str:
    value = filename.strip().replace("\\", "/").split("/")[-1]
    return (value or "lore-file")[:240]


def validate_lore_upload_contract(
    *,
    extension: str,
    content_type: str | None,
    file_bytes: bytes,
) -> None:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    allowed_content_types = {
        ".pdf": {"application/pdf"},
        ".docx": {DOCX_CONTENT_TYPE},
        ".txt": {"text/plain"},
        ".md": {"text/markdown", "text/plain", "text/x-markdown"},
    }
    if normalized_content_type not in allowed_content_types.get(extension, set()):
        raise CharacterLoreValidationError(
            "파일 확장자와 Content-Type이 일치하지 않습니다."
        )
    if not file_bytes:
        raise CharacterLoreValidationError("빈 설정집 파일은 업로드할 수 없습니다.")
    if extension == ".pdf":
        if not file_bytes.startswith(b"%PDF-"):
            raise CharacterLoreValidationError("유효한 PDF 파일이 아닙니다.")
        return
    if extension == ".docx":
        if not file_bytes.startswith(b"PK\x03\x04"):
            raise CharacterLoreValidationError("유효한 DOCX 파일이 아닙니다.")
        _preflight_docx(file_bytes)
        return
    if b"\x00" in file_bytes or file_bytes.startswith((b"%PDF-", b"PK\x03\x04")):
        raise CharacterLoreValidationError("텍스트 파일에서 바이너리 데이터가 감지됐습니다.")
    control_bytes = sum(
        byte < 32 and byte not in {9, 10, 13}
        for byte in file_bytes[: min(len(file_bytes), 64 * 1024)]
    )
    if control_bytes > max(1, min(len(file_bytes), 64 * 1024) // 100):
        raise CharacterLoreValidationError("유효한 텍스트 파일이 아닙니다.")


def _preflight_docx(file_bytes: bytes) -> None:
    try:
        with ZipFile(BytesIO(file_bytes)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise CharacterLoreValidationError("DOCX 내부 파일 수가 너무 많습니다.")
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise CharacterLoreValidationError("DOCX 필수 문서 구조가 없습니다.")

            total_uncompressed = 0
            total_xml = 0
            for entry in entries:
                path = PurePosixPath(entry.filename.replace("\\", "/"))
                if (
                    not path.parts
                    or path.is_absolute()
                    or ".." in path.parts
                    or ":" in path.parts[0]
                    or "\x00" in entry.filename
                ):
                    raise CharacterLoreValidationError(
                        "DOCX 내부 경로가 안전하지 않습니다."
                    )
                if entry.flag_bits & 0x1:
                    raise CharacterLoreValidationError(
                        "암호화된 DOCX 항목은 지원하지 않습니다."
                    )
                mode = entry.external_attr >> 16
                if stat.S_ISLNK(mode):
                    raise CharacterLoreValidationError(
                        "DOCX 내부 심볼릭 링크는 허용하지 않습니다."
                    )
                if entry.file_size > MAX_DOCX_ENTRY_BYTES:
                    raise CharacterLoreValidationError(
                        "DOCX 내부 파일 크기가 제한을 초과했습니다."
                    )
                if entry.file_size and (
                    entry.compress_size <= 0
                    or entry.file_size
                    > entry.compress_size * MAX_DOCX_COMPRESSION_RATIO
                ):
                    raise CharacterLoreValidationError(
                        "DOCX 압축 비율이 안전 한도를 초과했습니다."
                    )
                total_uncompressed += entry.file_size
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise CharacterLoreValidationError(
                        "DOCX 압축 해제 크기가 제한을 초과했습니다."
                    )

                if entry.filename.lower().endswith((".xml", ".rels")):
                    total_xml += entry.file_size
                    if total_xml > MAX_DOCX_XML_BYTES:
                        raise CharacterLoreValidationError(
                            "DOCX XML 크기가 제한을 초과했습니다."
                        )
                    xml_bytes = archive.read(entry)
                    lowered = xml_bytes.lower()
                    if b"<!doctype" in lowered or b"<!entity" in lowered:
                        raise CharacterLoreValidationError(
                            "DOCX 외부 엔터티 선언은 허용하지 않습니다."
                        )
    except CharacterLoreValidationError:
        raise
    except (BadZipFile, OSError, RuntimeError, ValueError) as exc:
        raise CharacterLoreValidationError("유효한 DOCX 파일이 아닙니다.") from exc


def _extract_text(
    extension: str,
    file_bytes: bytes,
    *,
    content_type: str | None = None,
) -> str:
    if content_type is not None:
        validate_lore_upload_contract(
            extension=extension,
            content_type=content_type,
            file_bytes=file_bytes,
        )
    if extension in {".txt", ".md"}:
        text = _decode_text_file(file_bytes)
    elif extension == ".pdf":
        text = _extract_pdf_text(file_bytes)
    elif extension == ".docx":
        text = _extract_docx_text(file_bytes)
    else:
        raise CharacterLoreValidationError("지원하지 않는 설정집 파일 형식입니다.")
    normalized = _normalize_text(text)
    if not normalized:
        raise CharacterLoreValidationError("설정집에서 텍스트를 추출하지 못했습니다.")
    return normalized


def _decode_text_file(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError:
        return file_bytes.decode("cp949", errors="replace")


def _extract_pdf_text(file_bytes: bytes) -> str:
    return _run_document_parser(".pdf", file_bytes)


def _extract_docx_text(file_bytes: bytes) -> str:
    return _run_document_parser(".docx", file_bytes)


def _run_document_parser(extension: str, file_bytes: bytes) -> str:
    result = _run_worker_process(
        _document_parser_worker,
        (extension, file_bytes),
        timeout_seconds=LORE_PARSER_TIMEOUT_SECONDS,
    )
    if not isinstance(result, str):
        raise CharacterLoreValidationError("문서 파서가 잘못된 결과를 반환했습니다.")
    return result


def _run_worker_process(
    target: Callable[..., None],
    args: tuple[Any, ...],
    *,
    timeout_seconds: float,
) -> Any:
    context = multiprocessing.get_context("spawn")
    parent_connection, child_connection = context.Pipe(duplex=False)
    process = context.Process(
        target=target,
        args=(child_connection, *args),
        daemon=True,
    )
    try:
        process.start()
        child_connection.close()
        if not parent_connection.poll(timeout_seconds):
            raise CharacterLoreValidationError(
                "문서 처리 시간이 안전 제한을 초과했습니다."
            )
        try:
            status_name, payload = parent_connection.recv()
        except (EOFError, OSError, ValueError) as exc:
            raise CharacterLoreValidationError(
                "격리된 문서 파서가 비정상 종료됐습니다."
            ) from exc
        if status_name == "ok":
            return payload
        if status_name == "validation":
            raise CharacterLoreValidationError(str(payload)[:500])
        raise CharacterLoreValidationError("문서 파서가 파일을 처리하지 못했습니다.")
    finally:
        parent_connection.close()
        child_connection.close()
        if process.is_alive():
            process.terminate()
            process.join(timeout=1)
        if process.is_alive():
            process.kill()
            process.join(timeout=1)
        if process.pid is not None and not process.is_alive():
            process.join(timeout=0.1)


def _document_parser_worker(
    connection: Any,
    extension: str,
    file_bytes: bytes,
) -> None:
    try:
        _apply_parser_resource_limits()
        if extension == ".pdf":
            text = _extract_pdf_text_in_process(file_bytes)
        elif extension == ".docx":
            text = _extract_docx_text_in_process(file_bytes)
        else:
            raise CharacterLoreValidationError("지원하지 않는 문서 형식입니다.")
        connection.send(("ok", text))
    except CharacterLoreValidationError as exc:
        connection.send(("validation", str(exc)[:500]))
    except BaseException:
        connection.send(("error", "document_parser_failed"))
    finally:
        connection.close()


def _extract_pdf_text_in_process(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise CharacterLoreValidationError("PDF 페이지 수가 제한을 초과했습니다.")
        object_count = reader.trailer.get("/Size", 0)
        if isinstance(object_count, int) and object_count > MAX_PDF_OBJECTS:
            raise CharacterLoreValidationError("PDF 객체 수가 제한을 초과했습니다.")
        parts: list[str] = []
        extracted_chars = 0
        for page in reader.pages:
            part = page.extract_text() or ""
            extracted_chars += len(part)
            if extracted_chars > MAX_PARSED_TEXT_CHARS:
                raise CharacterLoreValidationError(
                    "PDF 추출 텍스트가 제한을 초과했습니다."
                )
            if part.strip():
                parts.append(part)
    except CharacterLoreValidationError:
        raise
    except Exception as exc:
        raise CharacterLoreValidationError("PDF 텍스트를 추출하지 못했습니다.") from exc
    text = "\n\n".join(part.strip() for part in parts if part.strip())
    if not text.strip():
        raise CharacterLoreValidationError("스캔 이미지 PDF/OCR은 v1에서 지원하지 않습니다.")
    return text


def _extract_docx_text_in_process(file_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise CharacterLoreValidationError("DOCX 텍스트를 추출하지 못했습니다.") from exc
    parts: list[str] = []
    extracted_chars = 0
    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue
        extracted_chars += len(paragraph.text)
        if extracted_chars > MAX_PARSED_TEXT_CHARS:
            raise CharacterLoreValidationError(
                "DOCX 추출 텍스트가 제한을 초과했습니다."
            )
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                row_text = " | ".join(cells)
                extracted_chars += len(row_text)
                if extracted_chars > MAX_PARSED_TEXT_CHARS:
                    raise CharacterLoreValidationError(
                        "DOCX 추출 텍스트가 제한을 초과했습니다."
                    )
                parts.append(row_text)
    return "\n\n".join(parts)


_WINDOWS_PARSER_JOB_HANDLE: int | None = None


def _apply_parser_resource_limits() -> None:
    if os.name == "nt":
        _apply_windows_parser_resource_limits()
        return
    if os.name == "posix":
        import resource

        resource.setrlimit(
            resource.RLIMIT_AS,
            (LORE_PARSER_MEMORY_BYTES, LORE_PARSER_MEMORY_BYTES),
        )
        resource.setrlimit(
            resource.RLIMIT_CPU,
            (LORE_PARSER_CPU_SECONDS, LORE_PARSER_CPU_SECONDS),
        )
        return
    raise RuntimeError(f"Unsupported parser isolation platform: {sys.platform}")


def _apply_windows_parser_resource_limits() -> None:
    import ctypes
    from ctypes import wintypes

    class IoCounters(ctypes.Structure):
        _fields_ = [
            ("ReadOperationCount", ctypes.c_ulonglong),
            ("WriteOperationCount", ctypes.c_ulonglong),
            ("OtherOperationCount", ctypes.c_ulonglong),
            ("ReadTransferCount", ctypes.c_ulonglong),
            ("WriteTransferCount", ctypes.c_ulonglong),
            ("OtherTransferCount", ctypes.c_ulonglong),
        ]

    class BasicLimitInformation(ctypes.Structure):
        _fields_ = [
            ("PerProcessUserTimeLimit", ctypes.c_longlong),
            ("PerJobUserTimeLimit", ctypes.c_longlong),
            ("LimitFlags", wintypes.DWORD),
            ("MinimumWorkingSetSize", ctypes.c_size_t),
            ("MaximumWorkingSetSize", ctypes.c_size_t),
            ("ActiveProcessLimit", wintypes.DWORD),
            ("Affinity", ctypes.c_size_t),
            ("PriorityClass", wintypes.DWORD),
            ("SchedulingClass", wintypes.DWORD),
        ]

    class ExtendedLimitInformation(ctypes.Structure):
        _fields_ = [
            ("BasicLimitInformation", BasicLimitInformation),
            ("IoInfo", IoCounters),
            ("ProcessMemoryLimit", ctypes.c_size_t),
            ("JobMemoryLimit", ctypes.c_size_t),
            ("PeakProcessMemoryUsed", ctypes.c_size_t),
            ("PeakJobMemoryUsed", ctypes.c_size_t),
        ]

    kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
    kernel32.CreateJobObjectW.argtypes = [ctypes.c_void_p, wintypes.LPCWSTR]
    kernel32.CreateJobObjectW.restype = wintypes.HANDLE
    kernel32.SetInformationJobObject.argtypes = [
        wintypes.HANDLE,
        ctypes.c_int,
        ctypes.c_void_p,
        wintypes.DWORD,
    ]
    kernel32.SetInformationJobObject.restype = wintypes.BOOL
    kernel32.AssignProcessToJobObject.argtypes = [wintypes.HANDLE, wintypes.HANDLE]
    kernel32.AssignProcessToJobObject.restype = wintypes.BOOL
    kernel32.GetCurrentProcess.restype = wintypes.HANDLE

    job = kernel32.CreateJobObjectW(None, None)
    if not job:
        raise OSError(ctypes.get_last_error(), "CreateJobObjectW failed")
    information = ExtendedLimitInformation()
    information.BasicLimitInformation.PerProcessUserTimeLimit = (
        LORE_PARSER_CPU_SECONDS * 10_000_000
    )
    information.BasicLimitInformation.LimitFlags = 0x00000002 | 0x00000100 | 0x00002000
    information.ProcessMemoryLimit = LORE_PARSER_MEMORY_BYTES
    if not kernel32.SetInformationJobObject(
        job,
        9,
        ctypes.byref(information),
        ctypes.sizeof(information),
    ):
        error_code = ctypes.get_last_error()
        kernel32.CloseHandle(job)
        raise OSError(error_code, "SetInformationJobObject failed")
    if not kernel32.AssignProcessToJobObject(job, kernel32.GetCurrentProcess()):
        error_code = ctypes.get_last_error()
        kernel32.CloseHandle(job)
        raise OSError(error_code, "AssignProcessToJobObject failed")
    global _WINDOWS_PARSER_JOB_HANDLE
    _WINDOWS_PARSER_JOB_HANDLE = int(job)


def _normalize_text(text: str) -> str:
    value = text.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _split_lore_units(text: str) -> list[tuple[str | None, str]]:
    units: list[tuple[str | None, str]] = []
    section_hint: str | None = None
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            units.append((section_hint, " ".join(paragraph).strip()))
            paragraph = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        if _looks_like_section_heading(line):
            flush_paragraph()
            section_hint = _clean_heading(line)
            continue
        if _looks_like_boundary_line(line):
            flush_paragraph()
            units.append((section_hint, line))
            continue
        paragraph.append(line)
    flush_paragraph()
    return units


def _looks_like_section_heading(line: str) -> bool:
    stripped = line.strip().strip("#").strip()
    if not stripped or len(stripped) > 80:
        return False
    if line.lstrip().startswith("#"):
        return True
    if stripped.endswith(":"):
        return True
    return bool(re.match(r"^[\[(<【].+[\])>】]$", stripped))


def _clean_heading(line: str) -> str:
    return line.strip().strip("#").strip().rstrip(":")[:200]


def _looks_like_boundary_line(line: str) -> bool:
    return bool(
        re.match(r"^([-*•]|\d+[.)])\s+", line)
        or re.match(r"^(Q|A|문|답)\s*[:：]", line, re.IGNORECASE)
    )


def _split_long_unit(unit: str) -> list[str]:
    if len(unit) <= TARGET_CHUNK_MAX_CHARS:
        return [unit]
    sentences = re.split(r"(?<=[.!?。！？])\s+", unit)
    pieces: list[str] = []
    current = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if len(sentence) > TARGET_CHUNK_MAX_CHARS:
            if current:
                pieces.append(current)
                current = ""
            pieces.extend(
                sentence[index : index + TARGET_CHUNK_MAX_CHARS]
                for index in range(0, len(sentence), TARGET_CHUNK_MAX_CHARS)
            )
            continue
        if current and len(current) + len(sentence) + 1 > TARGET_CHUNK_MAX_CHARS:
            pieces.append(current)
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current:
        pieces.append(current)
    return pieces or [unit[:TARGET_CHUNK_MAX_CHARS]]


def _ensure_character_limits(
    db: Session,
    *,
    character_id: str,
    new_text_chars: int,
    new_chunk_count: int,
    replacing_text_chars: int = 0,
    replacing_chunk_count: int = 0,
) -> None:
    current_text_chars = max(0, _total_text_chars(db, character_id) - replacing_text_chars)
    current_chunk_count = max(0, _chunk_count(db, character_id) - replacing_chunk_count)
    if current_text_chars + new_text_chars > MAX_LORE_TEXT_CHARS_PER_CHARACTER:
        raise CharacterLoreValidationError("설정집 원문은 앵무당 최대 50,000자까지 가능합니다.")
    if current_chunk_count + new_chunk_count > MAX_LORE_CHUNKS_PER_CHARACTER:
        raise CharacterLoreValidationError("설정집 chunk는 앵무당 최대 100개까지 가능합니다.")


def _source_count(db: Session, character_id: str) -> int:
    return int(
        db.scalar(
            select(func.count(models.CharacterLoreSource.id)).where(
                models.CharacterLoreSource.character_id == character_id
            )
        )
        or 0
    )


def _total_text_chars(db: Session, character_id: str) -> int:
    return int(
        db.scalar(
            select(func.coalesce(func.sum(models.CharacterLoreSource.extracted_char_count), 0))
            .where(models.CharacterLoreSource.character_id == character_id)
        )
        or 0
    )


def _chunk_count(db: Session, character_id: str) -> int:
    return int(
        db.scalar(
            select(func.count(models.CharacterLoreChunk.id)).where(
                models.CharacterLoreChunk.character_id == character_id
            )
        )
        or 0
    )


def _ready_chunk_count(db: Session, character_id: str) -> int:
    return int(
        db.scalar(
            select(func.count(models.CharacterLoreChunk.id)).where(
                models.CharacterLoreChunk.character_id == character_id,
                models.CharacterLoreChunk.status == "ready",
            )
        )
        or 0
    )


def _status_read(db: Session, character_id: str) -> schemas.CharacterLoreStatusRead:
    source_count = _source_count(db, character_id)
    ready_source_count = int(
        db.scalar(
            select(func.count(models.CharacterLoreSource.id)).where(
                models.CharacterLoreSource.character_id == character_id,
                models.CharacterLoreSource.status.in_(("ready", "partial")),
            )
        )
        or 0
    )
    chunk_count = _chunk_count(db, character_id)
    ready_chunk_count = _ready_chunk_count(db, character_id)
    return schemas.CharacterLoreStatusRead(
        character_id=character_id,
        source_count=source_count,
        ready_source_count=ready_source_count,
        chunk_count=chunk_count,
        ready_chunk_count=ready_chunk_count,
        max_sources=MAX_LORE_SOURCES_PER_CHARACTER,
        max_text_chars=MAX_LORE_TEXT_CHARS_PER_CHARACTER,
        max_chunks=MAX_LORE_CHUNKS_PER_CHARACTER,
        max_file_bytes=MAX_LORE_FILE_BYTES,
    )


def _google_embedding_credential_for_character(
    db: Session, character_id: str
) -> _GoogleEmbeddingCredential:
    credential = agent_crud.get_character_credential(db, character_id)
    try:
        material = CredentialResolver.resolve_llm_credential(
            credential,
            purpose=CredentialPurpose.LORE_EMBEDDING,
            character_id=character_id,
        )
        if material.provider != "google":
            raise CredentialResolutionError("credential provider is not Google")
        api_key = material.reveal()
    except CredentialResolutionError as exc:
        raise CharacterLoreEmbeddingError("Google API key could not be decrypted.") from exc
    return _GoogleEmbeddingCredential(
        api_key=api_key,
        credential_id=credential.id,
        key_fingerprint=credential.key_fingerprint,
        provider=credential.provider,
    )


def _google_api_key_for_character(db: Session, character_id: str) -> str:
    credential = agent_crud.get_character_credential(db, character_id)
    try:
        material = CredentialResolver.resolve_llm_credential(
            credential,
            purpose=CredentialPurpose.LORE_EMBEDDING,
            character_id=character_id,
        )
        if material.provider != "google":
            raise CredentialResolutionError("credential provider is not Google")
        return material.reveal()
    except CredentialResolutionError as exc:
        raise CharacterLoreEmbeddingError("Google API key를 복호화하지 못했습니다.") from exc


def _existing_embeddings_by_hash(db: Session, character_id: str) -> dict[str, list[float]]:
    rows = db.scalars(
        select(models.CharacterLoreChunk).where(
            models.CharacterLoreChunk.character_id == character_id,
            models.CharacterLoreChunk.status == "ready",
            models.CharacterLoreChunk.embedding.is_not(None),
            models.CharacterLoreChunk.embedding_model == EMBEDDING_MODEL,
            models.CharacterLoreChunk.embedding_dimension == EMBEDDING_DIMENSION,
        )
    )
    result: dict[str, list[float]] = {}
    for chunk in rows:
        if chunk.content_hash not in result and chunk.embedding is not None:
            result[chunk.content_hash] = list(chunk.embedding)
    return result


def _chunk_embedding_input(draft: LoreChunkDraft) -> str:
    return f"title: {draft.section_hint or 'none'} | text: {draft.text}"


def _query_embedding_input(query: str) -> str:
    return f"task: search result | query: {query}"


def _embed_text(api_key: str, text: str) -> list[float]:
    adapter = get_embedding_adapter("google", EMBEDDING_MODEL)
    try:
        return adapter.embed_sync(
            EmbeddingRequest(
                api_key=api_key,
                model=EMBEDDING_MODEL,
                text=text,
                output_dimension=EMBEDDING_DIMENSION,
            )
        )
    except Exception as exc:
        safe_error = adapter.normalize_error(exc, api_key=api_key)
        raise CharacterLoreEmbeddingError(str(safe_error)) from exc


async def _embed_text_tracked(
    api_key: str,
    text: str,
    *,
    context: DirectLlmCallContext,
    tracker: RunLlmTracker,
) -> list[float]:
    await wait_for_provider_rate_limit(
        context=context,
        tracker=tracker,
        call_type="embed_content",
    )
    provider_call_order = tracker.next_provider_call_order()
    started = time.perf_counter()
    try:
        values = await asyncio.to_thread(_embed_text, api_key, text)
    except Exception as exc:
        tracker.record_embedding_call(
            context=context,
            provider_call_order=provider_call_order,
            status="error",
            duration_ms=int((time.perf_counter() - started) * 1000),
            failure_class=type(exc).__name__,
        )
        raise
    tracker.record_embedding_call(
        context=context,
        provider_call_order=provider_call_order,
        status="ok",
        duration_ms=int((time.perf_counter() - started) * 1000),
    )
    return values


def _format_recent_lore_usage(db: Session, *, character_id: str) -> str:
    chunks = list(
        db.scalars(
            select(models.CharacterLoreChunk)
            .where(
                models.CharacterLoreChunk.character_id == character_id,
                models.CharacterLoreChunk.last_used_at.is_not(None),
            )
            .order_by(
                models.CharacterLoreChunk.last_used_at.desc(),
                models.CharacterLoreChunk.id.desc(),
            )
            .limit(RECENT_LORE_USAGE_LIMIT)
        )
    )
    if not chunks:
        return "- none"
    lines = []
    for chunk in chunks:
        source_name = chunk.source.filename if chunk.source else "-"
        lines.append(
            f"- chunk_id={chunk.id}, source={source_name}, section={chunk.section_hint or '-'}, "
            f"last_used_at={chunk.last_used_at.isoformat() if chunk.last_used_at else '-'}, "
            f"usage_count={chunk.usage_count}"
        )
    return "\n".join(lines)


def _rerank_lore_candidates(rows: list[tuple[models.CharacterLoreChunk, float]]) -> list[RetrievedLoreChunk]:
    now = datetime.now(UTC)
    scored: list[tuple[float, models.CharacterLoreChunk, float]] = []
    for chunk, raw_distance in rows:
        distance = float(raw_distance or 0.0)
        penalty = min(0.2, max(0, chunk.usage_count) * 0.02)
        if chunk.last_used_at is not None:
            age = now - chunk.last_used_at
            if age <= RECENT_LORE_STRONG_PENALTY_WINDOW:
                penalty += 0.35
            elif age <= RECENT_LORE_SOFT_PENALTY_WINDOW:
                penalty += 0.18
        scored.append((distance + penalty, chunk, distance))
    scored.sort(key=lambda item: (item[0], item[1].usage_count, item[1].chunk_index))

    selected: list[RetrievedLoreChunk] = []
    used_sources: set[str] = set()
    used_sections: set[tuple[str, str]] = set()

    def append_candidate(chunk: models.CharacterLoreChunk, distance: float) -> None:
        selected.append(
            RetrievedLoreChunk(
                id=chunk.id,
                source_id=chunk.source_id,
                source_filename=chunk.source.filename if chunk.source else "-",
                section_hint=chunk.section_hint,
                text=chunk.text,
                distance=distance,
            )
        )
        used_sources.add(chunk.source_id)
        if chunk.section_hint:
            used_sections.add((chunk.source_id, chunk.section_hint))

    for _, chunk, distance in scored:
        if len(selected) >= RETRIEVAL_FINAL_LIMIT:
            break
        if chunk.source_id in used_sources and len(selected) < 2:
            continue
        if chunk.section_hint and (chunk.source_id, chunk.section_hint) in used_sections:
            continue
        append_candidate(chunk, distance)

    if len(selected) < min(3, len(scored)):
        selected_ids = {chunk.id for chunk in selected}
        for _, chunk, distance in scored:
            if len(selected) >= RETRIEVAL_FINAL_LIMIT:
                break
            if chunk.id in selected_ids:
                continue
            append_candidate(chunk, distance)
            selected_ids.add(chunk.id)
    return selected


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()
